import streamlit as st
from sentence_transformers import SentenceTransformer, util
from openai import OpenAI
import pdfplumber
import docx
import difflib
import re
import pandas as pd

# Load model
@st.cache_resource(show_spinner=True)
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

def read_pdf_pages(file):
    pages = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            pages.append(text if text else "")
    return pages

def read_docx_pages(file):
    doc = docx.Document(file)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    return ["\n".join(paragraphs)]  # Treat whole doc as one page

def read_txt_pages(file):
    text = file.getvalue().decode("utf-8")
    return [text]  # Treat whole txt as one page

def read_file_pages(file):
    if file.type == "application/pdf":
        return read_pdf_pages(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx_pages(file)
    else:
        return read_txt_pages(file)

def embed_text(embedder, texts):
    return embedder.encode(texts, convert_to_tensor=True)

def semantic_search(source_embedding, target_embeddings, top_k=1):
    hits = util.semantic_search(source_embedding, target_embeddings, top_k=top_k)
    return hits[0]

def openai_quality_assessment(client, source_text, translation_text):
    prompt = f"""
You are a translation quality evaluator.
Source page text: "{source_text}"
Translation page text: "{translation_text}"
Please provide a brief evaluation of the translation quality, highlighting any errors or issues.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",  # or "gpt-3.5-turbo"
        messages=[{"role": "user", "content": prompt}],
        max_tokens=150,
        temperature=0.5,
    )
    return response.choices[0].message["content"]

def highlight_glossary_terms(text, terms, color="#ADD8E6"):
    """
    Highlight all occurrences of glossary terms in the text with the given color.
    Case-insensitive whole word matching.
    """
    def replacer(match):
        return f'<span style="background-color:{color};">{match.group(0)}</span>'

    # Sort terms by length descending to avoid partial overlapping matches
    sorted_terms = sorted(terms, key=len, reverse=True)
    for term in sorted_terms:
        # Use regex for whole word, case-insensitive
        pattern = re.compile(r'\b' + re.escape(term) + r'\b', flags=re.IGNORECASE)
        text = pattern.sub(replacer, text)
    return text

def highlight_differences(text1, text2):
    """
    Highlights similar terms in green and different terms in red.
    Uses difflib.SequenceMatcher on word tokens.
    Returns HTML string with colored spans.
    """
    def tokenize(text):
        return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

    tokens1 = tokenize(text1)
    tokens2 = tokenize(text2)

    matcher = difflib.SequenceMatcher(None, tokens1, tokens2)
    highlighted_text = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for token in tokens1[i1:i2]:
                highlighted_text.append(f'<span style="background-color:#d4fcdc">{token}</span>')
        elif tag == 'replace' or tag == 'delete':
            for token in tokens1[i1:i2]:
                highlighted_text.append(f'<span style="background-color:#fcdcdc">{token}</span>')
        elif tag == 'insert':
            # Insertions in translation ignored here
            pass
        # Add space after each token except punctuation
        if i2 > i1:
            last_token = tokens1[i2-1]
            if re.match(r"\w", last_token):
                highlighted_text.append(" ")

    return "".join(highlighted_text).strip()

def main():
    st.title("Page-Level Translation Quality Checker with Glossary Highlighting")

    openai_api_key = st.text_input("Enter API key", type="password")
    client = None
    if openai_api_key:
        client = OpenAI(api_key=openai_api_key)

    embedder = load_model()

    source_file = st.file_uploader("Upload Benchmark document (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
    if source_file:
        source_pages = read_file_pages(source_file)
        st.success(f"Benchmark document loaded with {len(source_pages)} pages.")

        translation_files = st.file_uploader("Upload translation documents (multiple allowed)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        glossary_file = st.file_uploader("Upload optional glossary Excel file (with 'source' and 'translation' columns)", type=["xlsx"])

        glossary_source_terms = []
        glossary_translation_terms = []

        if glossary_file:
            try:
                df_glossary = pd.read_excel(glossary_file)
                if 'source' in df_glossary.columns and 'translation' in df_glossary.columns:
                    glossary_source_terms = df_glossary['source'].dropna().astype(str).tolist()
                    glossary_translation_terms = df_glossary['translation'].dropna().astype(str).tolist()
                    st.success(f"Loaded glossary with {len(glossary_source_terms)} source terms and {len(glossary_translation_terms)} translation terms.")
                else:
                    st.error("Glossary file must contain 'source' and 'translation' columns.")
            except Exception as e:
                st.error(f"Failed to read glossary file: {e}")

        if translation_files:
            translations = {}
            for file in translation_files:
                pages = read_file_pages(file)
                translations[file.name] = pages
            st.success(f"Loaded {len(translations)} translation documents.")

            page_index = st.number_input(f"Select benchmark page number (1 to {len(source_pages)})", min_value=1, max_value=len(source_pages), value=1)
            source_page_text = source_pages[page_index - 1]

            source_emb = embed_text(embedder, [source_page_text])

            for name, pages in translations.items():
                st.markdown(f"### Translation: {name}")
                translation_emb = embed_text(embedder, pages)
                hits = semantic_search(source_emb, translation_emb, top_k=1)
                best_hit = hits[0]
                best_page_text = pages[best_hit['corpus_id']]
                similarity = best_hit['score']

                st.write(f"Most similar page (similarity score: {similarity:.3f}):")

                # Highlight differences first
                highlighted_source = highlight_differences(source_page_text, best_page_text)
                highlighted_translation = highlight_differences(best_page_text, source_page_text)

                # Then highlight glossary terms in blue on top of difference highlights
                if glossary_source_terms:
                    highlighted_source = highlight_glossary_terms(highlighted_source, glossary_source_terms, color="#ADD8E6")
                if glossary_translation_terms:
                    highlighted_translation = highlight_glossary_terms(highlighted_translation, glossary_translation_terms, color="#ADD8E6")

                col1, col2 = st.columns(2)

                with col1:
                    st.markdown(f"### Benchmark page [{page_index}]:")
                    st.markdown(highlighted_source, unsafe_allow_html=True)

                with col2:
                    st.markdown(f"### Translation page [{best_hit['corpus_id'] + 1}]:")
                    st.markdown(highlighted_translation, unsafe_allow_html=True)

                if client:
                    with st.spinner("Evaluating translation quality..."):
                        assessment = openai_quality_assessment(client, source_page_text, best_page_text)
                    st.markdown("**Quality Assessment:**")
                    st.write(assessment)
                else:
                    st.info("OpenAI API key not provided. Showing similarity scores only.")

if __name__ == "__main__":
    main()
import streamlit as st
from sentence_transformers import SentenceTransformer, util
from openai import OpenAI
import pdfplumber
import docx
import difflib
import re

# Load model
@st.cache_resource(show_spinner=True)
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

def read_pdf_pages(file):
    pages = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            pages.append(text if text else "")
    return pages

def read_docx_pages(file):
    doc = docx.Document(file)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    return ["\n".join(paragraphs)]  # Treat whole doc as one page

def read_txt_pages(file):
    text = file.getvalue().decode("utf-8")
    return [text]  # Treat whole txt as one page

def read_file_pages(file):
    if file.type == "application/pdf":
        return read_pdf_pages(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx_pages(file)
    else:
        return read_txt_pages(file)

def embed_text(embedder, texts):
    return embedder.encode(texts, convert_to_tensor=True)

def semantic_search(source_embedding, target_embeddings, top_k=1):
    hits = util.semantic_search(source_embedding, target_embeddings, top_k=top_k)
    return hits[0]

def openai_quality_assessment(client, source_text, translation_text):
    prompt = f"""
You are a translation quality evaluator.
Source page text: "{source_text}"
Translation page text: "{translation_text}"
Please provide a brief evaluation of the translation quality, highlighting any errors or issues.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",  # or "gpt-3.5-turbo"
        messages=[{"role": "user", "content": prompt}],
        max_tokens=150,
        temperature=0.5, 
)
    # Updated for new API structure
    return response.choices[0].message["content"]

def highlight_differences(text1, text2):
    """
    Highlights similar terms in green and different terms in red.
    Uses difflib.SequenceMatcher on word tokens.
    Returns HTML string with colored spans.
    """
    def tokenize(text):
        return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

    tokens1 = tokenize(text1)
    tokens2 = tokenize(text2)

    matcher = difflib.SequenceMatcher(None, tokens1, tokens2)
    highlighted_text = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for token in tokens1[i1:i2]:
                highlighted_text.append(f'<span style="background-color:#d4fcdc">{token}</span>')
        elif tag == 'replace' or tag == 'delete':
            for token in tokens1[i1:i2]:
                highlighted_text.append(f'<span style="background-color:#fcdcdc">{token}</span>')
        elif tag == 'insert':
            # Insertions in translation ignored here
            pass
        # Add space after each token except punctuation
        if i2 > i1:
            last_token = tokens1[i2-1]
            if re.match(r"\w", last_token):
                highlighted_text.append(" ")

    return "".join(highlighted_text).strip()

def main():
    st.title("Page-Level Translation Quality Checker")

    openai_api_key = st.text_input("Enter API key", type="password")
    client = None
    if openai_api_key:
        client = OpenAI(api_key=openai_api_key)

    embedder = load_model()

    source_file = st.file_uploader("Upload Benchmark document (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
    if source_file:
        source_pages = read_file_pages(source_file)
        st.success(f"Benchmark document loaded with {len(source_pages)} pages.")

        translation_files = st.file_uploader("Upload translation documents (multiple allowed)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        if translation_files:
            translations = {}
            for file in translation_files:
                pages = read_file_pages(file)
                translations[file.name] = pages
            st.success(f"Loaded {len(translations)} translation documents.")

            page_index = st.number_input(f"Select benchmark page number (1 to {len(source_pages)})", min_value=1, max_value=len(source_pages), value=1)
            source_page_text = source_pages[page_index - 1]

            source_emb = embed_text(embedder, [source_page_text])

            for name, pages in translations.items():
                st.markdown(f"### Translation: {name}")
                translation_emb = embed_text(embedder, pages)
                hits = semantic_search(source_emb, translation_emb, top_k=1)
                best_hit = hits[0]
                best_page_text = pages[best_hit['corpus_id']]
                similarity = best_hit['score']

                st.write(f"Most similar page (similarity score: {similarity:.3f}):")

                highlighted_source = highlight_differences(source_page_text, best_page_text)
                highlighted_translation = highlight_differences(best_page_text, source_page_text)

                col1, col2 = st.columns(2)

                with col1:
                    st.markdown(f"### Benchmark page [{page_index}]:")
                    st.markdown(highlighted_source, unsafe_allow_html=True)

                with col2:
                    st.markdown(f"### Translation page [{best_hit['corpus_id'] + 1}]:")
                    st.markdown(highlighted_translation, unsafe_allow_html=True)

                if client:
                    with st.spinner("Evaluating translation quality..."):
                        assessment = openai_quality_assessment(client, source_page_text, best_page_text)
                    st.markdown("**Quality Assessment:**")
                    st.write(assessment)
                else:
                    st.info("OpenAI API key not provided. Showing similarity scores only.")

if __name__ == "__main__":
    main()
